import tkinter as tk
from tkinter import ttk, messagebox
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from datetime import datetime
from docx import Document
from fpdf import FPDF

class JobScraperApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Job Scraper")
        self.root.geometry("1200x600")  # Increase window size to accommodate more columns

        # Load the Azure theme
        self.root.tk.call("source", "azure.tcl")
        self.current_theme = "dark"
        self.root.tk.call("set_theme", self.current_theme)

        # Define URL selectors for different websites
        self.selectors = {
            'https://example.com/jobs': {
                'title': '.job-title',
                'company': '.company-name',
                'date': '.date-posted',
                'link': '.apply-link'
            },
            'https://anotherexample.com/careers': {
                'title': '.post-title',
                'company': '.org-name',
                'date': '.posted-date',
                'link': '.job-link'
            }
        }

        # Setup UI components
        self.setup_ui()

    def setup_ui(self):
        # Function to toggle theme
        def toggle_theme():
            if self.current_theme == "dark":
                self.current_theme = "light"
            else:
                self.current_theme = "dark"
            self.root.tk.call("set_theme", self.current_theme)

        # Function to export data to Word and PDF
        def export_data():
            if not self.tree.get_children():
                messagebox.showwarning("Export Error", "No data to export.")
                return

            # Export to Word
            doc = Document()
            doc.add_heading('Job Postings', level=1)

            table = doc.add_table(rows=1, cols=len(self.columns))
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(self.columns):
                hdr_cells[i].text = column

            for item in self.tree.get_children():
                row_data = self.tree.item(item, 'values')
                row_cells = table.add_row().cells
                for i, value in enumerate(row_data):
                    row_cells[i].text = value

            doc.save('job_postings.docx')

            # Export to PDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            col_width = pdf.get_string_width(max(self.columns, key=len)) + 10
            row_height = 10

            # Add header
            for column in self.columns:
                pdf.cell(col_width, row_height, column, border=1)
            pdf.ln(row_height)

            # Add rows
            for item in self.tree.get_children():
                row_data = self.tree.item(item, 'values')
                for value in row_data:
                    pdf.cell(col_width, row_height, str(value), border=1)
                pdf.ln(row_height)

            pdf.output('job_postings.pdf')

            messagebox.showinfo("Export Successful", "Data exported to job_postings.docx and job_postings.pdf")

        # Define UI components
        self.url_label = ttk.Label(self.root, text="Enter Job Listing URL:")
        self.url_label.pack(pady=10)

        self.url_entry = ttk.Entry(self.root, width=80)
        self.url_entry.pack(pady=5)

        self.fetch_button = ttk.Button(self.root, text="Fetch Job Postings", command=self.fetch_job_postings_selenium)
        self.fetch_button.pack(pady=10)

        self.clear_button = ttk.Button(self.root, text="Clear", command=self.clear_treeview)
        self.clear_button.pack(pady=5)

        self.toggle_theme_button = ttk.Button(self.root, text="Toggle Theme", command=toggle_theme)
        self.toggle_theme_button.pack(pady=5)

        self.export_button = ttk.Button(self.root, text="Export Data", command=export_data)
        self.export_button.pack(pady=5)

        # Define columns for Treeview
        self.columns = ('Title', 'Company', 'Date Posted', 'Job Link')

        # Treeview for displaying job listings
        self.tree = ttk.Treeview(self.root, columns=self.columns, show='headings', height=15)
        self.tree.heading('Title', text='Title')
        self.tree.heading('Company', text='Company')
        self.tree.heading('Date Posted', text='Date Posted')
        self.tree.heading('Job Link', text='Job Link')

        self.tree.column('Title', width=200)
        self.tree.column('Company', width=150)
        self.tree.column('Date Posted', width=100)
        self.tree.column('Job Link', width=250)

        self.tree.pack(pady=20)

        # Status label
        self.status_label = ttk.Label(self.root, text="")
        self.status_label.pack(pady=10)

    # Function to fetch job postings using Selenium
    def fetch_job_postings_selenium(self):
        options = Options()
        options.headless = False  # Set this to False to keep Chrome open
        options.add_argument('--ignore-certificate-errors')  # Ignore SSL certificate errors
        options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3")

        driver_path = r'C:\Users\M\Downloads\Compressed\chromedriver-win64\chromedriver.exe'
        service = Service(driver_path)
        driver = webdriver.Chrome(service=service, options=options)

        all_job_listings = []
        url = self.url_entry.get()  # Get the URL from the entry box

        if not url or url not in self.selectors:
            messagebox.showerror("Input Error", "Unsupported URL or No selectors defined.")
            return

        selectors = self.selectors.get(url)

        try:
            driver.get(url)  # Open the URL
            print(f"Fetching URL: {url}")

            # Wait for the page to load
            WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.TAG_NAME, 'body'))
            )

            # Print the page source for debugging
            print(driver.page_source)

            # Wait for job listings to load
            job_listings = WebDriverWait(driver, 30).until(
                EC.presence_of_all_elements_located((By.CSS_SELECTOR, '.job-listing'))  # Adjust selector as needed
            )

            if not job_listings:
                self.status_label.config(text="No job listings found.")
                print("No job listings found.")
                return

            for job in job_listings:
                try:
                    title = job.find_element(By.CSS_SELECTOR, selectors['title']).text
                    company = job.find_element(By.CSS_SELECTOR, selectors['company']).text
                    date_posted_str = job.find_element(By.CSS_SELECTOR, selectors['date']).text
                    job_link = job.find_element(By.CSS_SELECTOR, selectors['link']).get_attribute('href')

                    # Convert date_posted_str to datetime object
                    try:
                        date_posted = datetime.strptime(date_posted_str, '%d %b %Y')  # Adjust format as needed
                    except ValueError:
                        print(f"Error parsing date: {date_posted_str}")
                        continue

                    # Check if the job posting is from the current month and year
                    if date_posted.month == datetime.now().month and date_posted.year == datetime.now().year:
                        all_job_listings.append((title, company, date_posted_str, job_link))
                except Exception as e:
                    print(f"Error extracting job details: {e}")

        except Exception as e:
            self.status_label.config(text=f"Error fetching data with Selenium: {e}")
            print(f"Error fetching data with Selenium: {e}")
            return

        # Clear existing entries in the Treeview
        for item in self.tree.get_children():
            self.tree.delete(item)

        # Insert new job postings into the Treeview
        if not all_job_listings:
            self.status_label.config(text="No job postings found for the current month.")
            messagebox.showinfo("No Data", "No job postings found for the current month.")
            print("No job postings for the current month.")
        else:
            for job in all_job_listings:
                self.tree.insert('', tk.END, values=job)
            self.status_label.config(text="Job postings fetched successfully.")
            messagebox.showinfo("Fetch Successful", "Job postings for the current month fetched successfully.")
            print("Job postings successfully fetched.")

        # Close the browser after fetching data
        driver.quit()

    # Function to clear the Treeview and status label
    def clear_treeview(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        self.status_label.config(text="Treeview cleared.")

# Initialize main window
root = tk.Tk()
app = JobScraperApp(root)

# Run the main loop
root.mainloop()
